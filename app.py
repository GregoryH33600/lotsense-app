import streamlit as st
import ezdxf
import tempfile
import pandas as pd
import re
from docx import Document
from io import BytesIO

st.set_page_config(page_title="LotSense Pro", layout="wide")
st.title("🏢 LotSense Pro - DXF/DWG → EDD-RCP")

# Upload
uploaded_file = st.file_uploader("Upload DXF (DWG non converti à DXF)", type=["dxf"])
if uploaded_file:
    st.success(f"Fichier chargé : {uploaded_file.name}")

    # Sauvegarde temporaire
    with tempfile.NamedTemporaryFile(delete=False, suffix=".dxf") as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    # Lecture DXF
    doc = ezdxf.readfile(tmp_path)
    msp = doc.modelspace()

    lots = []
    for entity in msp:
        if entity.dxftype() == "TEXT":
            text = entity.dxf.text
            if "Lot" in text:
                # surface détectée
                surface_match = re.search(r"(\d+\.?\d*)\s*m²", text)
                surface = surface_match.group(1)+" m²" if surface_match else "Non détectée"
                # catégorie automatique partie commune
                category = "Partie privée"
                if any(k in text.lower() for k in ["hall", "escalier", "parking", "terrasse"]):
                    category = "Partie commune spéciale"
                lots.append({
                    "lot": text.replace("Lot","").split()[0],
                    "surface": surface,
                    "niveau": "Inconnu",
                    "catégorie": category
                })

    # Tableau affichage
    if lots:
        df = pd.DataFrame(lots)
        st.write("### Lots détectés")
        st.table(df)

        # Export Excel
        output_excel = BytesIO()
        df.to_excel(output_excel, index=False)
        st.download_button("📥 Télécharger Excel", data=output_excel,
                           file_name="lots_detectes.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # Génération EDD-RCP Word
        docx_file = Document()
        docx_file.add_heading('EDD-RCP Copropriété', 0)
        docx_file.add_paragraph("Liste des lots et surfaces :")
        table = docx_file.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Lot'
        hdr_cells[1].text = 'Surface'
        hdr_cells[2].text = 'Niveau'
        hdr_cells[3].text = 'Catégorie'
        for lot in lots:
            row_cells = table.add_row().cells
            row_cells[0].text = lot['lot']
            row_cells[1].text = lot['surface']
            row_cells[2].text = lot['niveau']
            row_cells[3].text = lot['catégorie']

        # Export Word
        output_docx = BytesIO()
        docx_file.save(output_docx)
        st.download_button("📥 Télécharger EDD-RCP Word", data=output_docx,
                           file_name="EDD-RCP.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
